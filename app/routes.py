from flask import Blueprint, render_template, request, jsonify, session, redirect, url_for, flash, send_file, Response
import io, csv
from datetime import datetime, timedelta

from sqlalchemy import func

from .models import db, Subject, Quiz, Question, Option, Teacher
from .models import Student, QuizAttempt, AttemptAnswer, Conversation, Message

main = Blueprint("main", __name__)


@main.route("/")
def select_role():
    return render_template("select_role.html")

@main.route("/student/dashboard")
def student_dashboard():
    # show available quizzes from all teachers (most recent first)
    from .models import Quiz, Teacher
    quizzes = Quiz.query.order_by(Quiz.created_at.desc()).limit(10).all()
    # eager load teacher info if available
    # template will handle missing teacher gracefully
    # try to fetch student's daily goal
    student_id = session.get('student_id')
    student = None
    daily_goal = 1
    completed_today = 0
    if student_id:
        student = Student.query.filter_by(id=student_id).first()
        if student:
            try:
                daily_goal = int(student.daily_goal or 1)
            except Exception:
                daily_goal = 1
        # compute completed today (simple count of attempts in last 24h)
        try:
            from datetime import datetime, timedelta
            now = datetime.utcnow()
            since = now - timedelta(days=1)
            completed_today = QuizAttempt.query.filter(QuizAttempt.student_id == student_id, QuizAttempt.completed_at != None, QuizAttempt.completed_at >= since).count()
        except Exception:
            completed_today = 0

    # compute total quizzes taken by this student
    quizzes_taken = 0
    if student_id:
        try:
            quizzes_taken = QuizAttempt.query.filter(QuizAttempt.student_id == student_id, QuizAttempt.completed_at != None).count()
        except Exception:
            quizzes_taken = 0

    # compute average percent score for this student across completed attempts
    avg_score = None
    if student_id:
        try:
            # use SQL aggregate to compute average of percent where percent is not null
            avg_val = db.session.query(func.avg(QuizAttempt.percent)).filter(QuizAttempt.student_id == student_id, QuizAttempt.completed_at != None, QuizAttempt.percent != None).scalar()
            if avg_val is not None:
                # round to 1 decimal place
                try:
                    avg_score = round(float(avg_val), 1)
                except Exception:
                    avg_score = None
        except Exception:
            avg_score = None

        # compute consecutive days streak based on completed quiz attempts
        days_streak = 0
        if student_id:
            try:
                # fetch distinct completed dates (UTC date portion)
                attempts = (
                    QuizAttempt.query.filter(QuizAttempt.student_id == student_id, QuizAttempt.completed_at != None)
                    .with_entities(QuizAttempt.completed_at)
                    .order_by(QuizAttempt.completed_at.desc())
                    .all()
                )
                # build a set of date objects (UTC)
                completed_dates = set()
                for (dt,) in attempts:
                    if dt is None:
                        continue
                    try:
                        # if it's a SQL expression (func.now()) it will be datetime when loaded
                        d = dt.date()
                    except Exception:
                        # fallback: attempt to parse or skip
                        continue
                    completed_dates.add(d)

                # compute streak ending today (UTC)
                today = datetime.utcnow().date()
                cur = today
                streak = 0
                # count backward while each consecutive date exists
                while cur in completed_dates:
                    streak += 1
                    cur = cur - timedelta(days=1)
                days_streak = streak
            except Exception:
                days_streak = 0

        # compute student's global rank (dense ranking) based on average percent across completed attempts
        current_student_rank = None
        current_student_avg = None
        try:
            averages = (
                db.session.query(
                    Student.id.label('student_id'),
                    func.avg(QuizAttempt.percent).label('avg_percent')
                )
                .join(QuizAttempt, QuizAttempt.student_id == Student.id)
                .filter(QuizAttempt.completed_at != None, QuizAttempt.percent != None)
                .group_by(Student.id)
                .order_by(func.avg(QuizAttempt.percent).desc())
                .all()
            )

            last_score = None
            dense_rank = 0
            for row in averages:
                avg = float(row.avg_percent) if row.avg_percent is not None else 0.0
                if last_score is None or avg != last_score:
                    dense_rank += 1
                last_score = avg
                if row.student_id == student_id:
                    current_student_rank = dense_rank
                    current_student_avg = round(avg, 1)
                    break
        except Exception:
            current_student_rank = None
            current_student_avg = None

        return render_template("student/dashboard.html", quizzes=quizzes, daily_goal=daily_goal, daily_completed=completed_today, quizzes_taken=quizzes_taken, avg_score=avg_score, days_streak=days_streak, current_student_rank=current_student_rank, current_student_avg=current_student_avg)


@main.route("/student/quizzes")
def student_quizzes():
    # show all available quizzes for students (most recent first)
    quizzes = Quiz.query.order_by(Quiz.created_at.desc()).all()
    return render_template("student/quizzes.html", quizzes=quizzes)

@main.route('/student/quizzes/take/<int:quiz_id>', methods=["GET", "POST"])
def student_take_quiz(quiz_id):
    # GET: render quiz for student to take
    if request.method == 'GET':
        quiz = Quiz.query.filter_by(id=quiz_id).first()
        if not quiz:
            flash('Quiz not found.', 'error')
            return redirect(url_for('main.student_quizzes'))
        # Render template with quiz and questions/options
        return render_template('student/take_quiz.html', quiz=quiz)

    # POST: grade submission (expect JSON payload with answers)
    # Payload format: { "answers": { "<question_id>": "<answer>", ... } }
    data = request.get_json() or {}
    answers = data.get('answers', {})
    # optional elapsed time (seconds) from client
    time_taken_seconds = data.get('time_taken_seconds')
    quiz = Quiz.query.filter_by(id=quiz_id).first()
    if not quiz:
        return jsonify({"error": "Quiz not found"}), 404

    total = len(quiz.questions)
    correct_count = 0
    details = []
    for q in quiz.questions:
        qid = str(q.id)
        given = answers.get(qid)
        is_correct = False
        # normalize True/False
        if q.type == 'tf':
            # Accept boolean or string
            if isinstance(given, bool):
                given_norm = 'True' if given else 'False'
            else:
                given_norm = str(given)
            is_correct = given_norm == q.correct_answer
        else:
            # mc: compare string keys (A/B/C...)
            if given is not None:
                is_correct = str(given).strip() == str(q.correct_answer).strip()

        if is_correct:
            correct_count += 1

        details.append({
            "question_id": q.id,
            "given": given,
            "correct_answer": q.correct_answer,
            "is_correct": is_correct,
        })

    score = correct_count
    percent = (score / total * 100) if total > 0 else 0
    result = {
        "quiz_id": quiz.id,
        "total_questions": total,
        "correct": score,
        "percent": percent,
        "details": details,
    }
    # attach human-friendly time if provided
    if isinstance(time_taken_seconds, int) or (isinstance(time_taken_seconds, float)):
        secs = int(time_taken_seconds)
        mm = secs // 60
        ss = secs % 60
        result['time_taken_seconds'] = secs
        result['time_taken'] = f"{mm:02d}:{ss:02d}"
    # store last result in session for display on results page
    # Persist attempt to DB if a student is logged in
    attempt_id = None
    student_id = session.get('student_id')
    if student_id:
        try:
            attempt = QuizAttempt(quiz_id=quiz.id, student_id=student_id, completed_at=db.func.now(), score=score, percent=percent, time_taken_seconds=time_taken_seconds)
            db.session.add(attempt)
            db.session.flush()  # get attempt.id
            # store per-question answers
            for d in details:
                aa = AttemptAnswer(attempt_id=attempt.id, question_id=d.get('question_id'), given_answer=str(d.get('given')) if d.get('given') is not None else None, is_correct=d.get('is_correct'))
                db.session.add(aa)
            db.session.commit()
            attempt_id = attempt.id
            # include attempt id in result for traceability
            result['attempt_id'] = attempt_id
        except Exception:
            db.session.rollback()
            # If DB save fails, continue gracefully and just keep result in session
            attempt_id = None

    try:
        # include attempt id in session result when available
        if attempt_id:
            result['_attempt_id'] = attempt_id
        session['last_quiz_result'] = result
    except Exception:
        # if session can't store (very large), just return result JSON
        return jsonify(result), 200

    return jsonify({"redirect": url_for('main.student_quiz_results')}), 200

@main.route("/student/quizzes/results")
def student_quiz_results():
    # display the most recent quiz result stored in session (one-off)
    result = session.pop('last_quiz_result', None)
    quiz = None
    counts = {
        'correct': 0,
        'incorrect': 0,
        'unanswered': 0,
    }
    if result:
        quiz_id = result.get('quiz_id')
        if quiz_id:
            quiz = Quiz.query.filter_by(id=quiz_id).first()
        details = result.get('details', []) or []
        total = result.get('total_questions', 0)
        correct = result.get('correct', 0)
        unanswered = 0
        for d in details:
            given = d.get('given')
            if given is None or (isinstance(given, str) and given.strip() == ''):
                unanswered += 1
        incorrect = max(0, total - correct - unanswered)
        counts['correct'] = correct
        counts['incorrect'] = incorrect
        counts['unanswered'] = unanswered

    return render_template("student/quiz_results.html", quiz=quiz, result=result, counts=counts)

@main.route("/teacher/dashboard")
def teacher_dashboard():
    teacher_id = session.get('teacher_id')
    recent_quizzes = []
    if teacher_id:
        recent_quizzes = (
            Quiz.query.filter_by(teacher_id=teacher_id)
            .order_by(Quiz.id.desc())
            .limit(5)
            .all()
        )
    # compute totals and recent deltas (last 7 days)
    try:
        now = datetime.utcnow()
        cutoff = now - timedelta(days=7)

        if teacher_id:
            total_quizzes = Quiz.query.filter_by(teacher_id=teacher_id).count()
            new_quizzes = (
                Quiz.query.filter(Quiz.teacher_id == teacher_id, Quiz.created_at >= cutoff).count()
            )

            # total_subjects as distinct subjects used by this teacher's quizzes
            total_subjects = (
                db.session.query(func.count(func.distinct(Quiz.subject_id)))
                .filter(Quiz.teacher_id == teacher_id, Quiz.subject_id != None)
                .scalar()
                or 0
            )

            # new_subjects: distinct subjects which had a quiz created in the cutoff window
            new_subjects = (
                db.session.query(func.count(func.distinct(Quiz.subject_id)))
                .filter(Quiz.teacher_id == teacher_id, Quiz.created_at >= cutoff, Quiz.subject_id != None)
                .scalar()
                or 0
            )
        else:
            total_quizzes = Quiz.query.count()
            new_quizzes = Quiz.query.filter(Quiz.created_at >= cutoff).count()
            total_subjects = Subject.query.count()
            new_subjects = (
                db.session.query(func.count(func.distinct(Quiz.subject_id)))
                .filter(Quiz.created_at >= cutoff, Quiz.subject_id != None)
                .scalar()
                or 0
            )

        # total students (all students in system)
        try:
            total_students = Student.query.count()
        except Exception:
            total_students = 0

        # new students: Student model has no created_at column; default to 0
        new_students = 0
    except Exception:
        total_quizzes = None
        new_quizzes = None
        total_subjects = None
        new_subjects = None
        total_students = None
        new_students = None

    return render_template(
        "teacher/dashboard.html",
        recent_quizzes=recent_quizzes,
        total_quizzes=total_quizzes,
        new_quizzes=new_quizzes,
        total_subjects=total_subjects,
        new_subjects=new_subjects,
        total_students=total_students,
        new_students=new_students,
    )


@main.route("/teacher/subjects")
def teacher_subjects():
    subjects = Subject.query.all()
    return render_template("teacher/subjects.html", subjects=subjects)


@main.route("/teacher/subjects/create", methods=["GET", "POST"])
def teacher_create_subject():
    if request.method == "GET":
        return render_template("teacher/create_subject.html")
    # POST: create subject from modal
    if request.is_json:
        data = request.get_json()
    else:
        data = request.form
    name = data.get('name')
    code = data.get('code')
    description = data.get('description')
    category = data.get('category')
    grade_level = data.get('grade_level')
    if not name:
        return jsonify({"error": "Subject name required"}), 400
    # Check for duplicate
    if Subject.query.filter_by(name=name).first():
        return jsonify({"error": "Subject already exists"}), 400
    subject = Subject(
        name=name,
        code=code,
        description=description,
        category=category,
        grade_level=grade_level
    )
    db.session.add(subject)
    db.session.commit()
    return jsonify({"status": "success", "subject_id": subject.id}), 201


@main.route("/teacher/quizzes")
def teacher_quizzes():
    from .models import Quiz
    teacher_id = session.get('teacher_id')
    quizzes = []
    if teacher_id:
        quizzes = Quiz.query.filter_by(teacher_id=teacher_id).order_by(Quiz.id.desc()).all()
    return render_template("teacher/quizzes.html", quizzes=quizzes)


@main.route("/teacher/quizzes/create", methods=["GET", "POST"])
def teacher_create_quiz():
    if request.method == "GET":
        subjects = Subject.query.all()
        return render_template("teacher/create_quiz.html", subjects=subjects)

    # POST: expect JSON payload describing quiz
    data = request.get_json()
    if not data:
        return jsonify({"error": "Invalid or missing JSON payload"}), 400

    title = data.get('title')
    subject_name = data.get('subject')
    time_limit = data.get('time_limit')
    difficulty = data.get('difficulty')
    description = data.get('description')
    questions = data.get('questions', [])

    if not title or not subject_name:
        return jsonify({"error": "Missing required fields: title and subject"}), 400

    # ensure a teacher is logged in
    teacher_id = session.get('teacher_id')
    if not teacher_id:
        return jsonify({"error": "Authentication required: teacher must be logged in"}), 401

    try:
        # get or create subject
        subject = Subject.query.filter_by(name=subject_name).first()
        if not subject:
            subject = Subject(name=subject_name)
            db.session.add(subject)
            db.session.flush()

        quiz = Quiz(title=title, description=description, time_limit=time_limit or 0, difficulty=difficulty, subject=subject, teacher_id=teacher_id)
        db.session.add(quiz)
        db.session.flush()

        for q in questions:
            qtype = q.get('type')
            qtext = q.get('text')
            correct = q.get('correct')
            opts = q.get('options', [])
            if not qtype or not qtext or correct is None:
                raise ValueError('Invalid question payload')

            question = Question(quiz=quiz, type=qtype, text=qtext, correct_answer=str(correct))
            db.session.add(question)
            db.session.flush()

            # options for MC questions
            if qtype == 'mc':
                for opt in opts:
                    key = opt.get('key')
                    text = opt.get('text')
                    if key and text is not None:
                        option = Option(question=question, key=key, text=text)
                        db.session.add(option)

        db.session.commit()
        return jsonify({"status": "success", "quiz_id": quiz.id}), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500


@main.route("/teacher/students")
def teacher_students():
    # Show list of students with total quizzes taken and average percent score
    try:
        # Use a left join aggregate to compute counts and averages per student in one query
        rows = (
            db.session.query(
                Student.id.label('id'),
                Student.name.label('name'),
                Student.email.label('email'),
                func.count(QuizAttempt.id).label('total_taken'),
                func.avg(QuizAttempt.percent).label('avg_percent')
            )
            .outerjoin(QuizAttempt, QuizAttempt.student_id == Student.id)
            .group_by(Student.id)
            .order_by(Student.name.asc())
            .all()
        )

        students = []
        for r in rows:
            avg = None
            try:
                if r.avg_percent is not None:
                    avg = round(float(r.avg_percent), 1)
            except Exception:
                avg = None

            students.append({
                'id': r.id,
                'name': r.name,
                'email': r.email,
                'total_taken': int(r.total_taken or 0),
                'avg_percent': avg,
            })
    except Exception:
        # On error, fallback to simple list of students without metrics
        students = []
        try:
            for s in Student.query.order_by(Student.name.asc()).all():
                students.append({'id': s.id, 'name': s.name, 'email': s.email, 'total_taken': 0, 'avg_percent': None})
        except Exception:
            students = []

    return render_template("teacher/students.html", students=students)


@main.route('/teacher/students/export')
def teacher_export_students():
    """Export the students list in TXT, CSV (Excel) or PDF format.
    Usage: /teacher/students/export?format=csv|txt|pdf
    """
    fmt = (request.args.get('format') or 'csv').lower()

    # build students list (same as teacher_students)
    try:
        rows = (
            db.session.query(
                Student.id.label('id'),
                Student.name.label('name'),
                Student.email.label('email'),
                func.count(QuizAttempt.id).label('total_taken'),
                func.avg(QuizAttempt.percent).label('avg_percent')
            )
            .outerjoin(QuizAttempt, QuizAttempt.student_id == Student.id)
            .group_by(Student.id)
            .order_by(Student.name.asc())
            .all()
        )

        students = []
        for r in rows:
            avg = None
            try:
                if r.avg_percent is not None:
                    avg = round(float(r.avg_percent), 1)
            except Exception:
                avg = None

            students.append({
                'id': r.id,
                'name': r.name,
                'email': r.email,
                'total_taken': int(r.total_taken or 0),
                'avg_percent': avg,
            })
    except Exception:
        # fallback to simple list
        students = []
        try:
            for s in Student.query.order_by(Student.name.asc()).all():
                students.append({'id': s.id, 'name': s.name, 'email': s.email, 'total_taken': 0, 'avg_percent': None})
        except Exception:
            students = []

    # TXT output
    if fmt == 'txt':
        out = io.StringIO()
        out.write('Name\tEmail\tQuizzes Taken\tAvg Score\n')
        for s in students:
            avg = '' if s.get('avg_percent') is None else f"{s.get('avg_percent')}%"
            out.write(f"{s.get('name')}\t{s.get('email')}\t{s.get('total_taken')}\t{avg}\n")
        bio = io.BytesIO()
        bio.write(out.getvalue().encode('utf-8'))
        bio.seek(0)
        return send_file(bio, as_attachment=True, download_name='students.txt', mimetype='text/plain')

    # CSV (Excel-friendly)
    if fmt == 'csv' or fmt == 'excel':
        out = io.StringIO()
        writer = csv.writer(out)
        writer.writerow(['Name', 'Email', 'Quizzes Taken', 'Avg Score'])
        for s in students:
            avg = '' if s.get('avg_percent') is None else str(s.get('avg_percent'))
            writer.writerow([s.get('name'), s.get('email'), s.get('total_taken'), avg])
        bio = io.BytesIO()
        bio.write(out.getvalue().encode('utf-8-sig'))
        bio.seek(0)
        return send_file(bio, as_attachment=True, download_name='students.csv', mimetype='text/csv')

    # PDF (optional dependency: reportlab)
    if fmt == 'pdf':
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
        except Exception:
            flash('PDF export requires the reportlab package to be installed.', 'error')
            return redirect(url_for('main.teacher_students'))

        bio = io.BytesIO()
        c = canvas.Canvas(bio, pagesize=letter)
        width, height = letter
        x = 40
        y = height - 40
        c.setFont('Helvetica-Bold', 14)
        c.drawString(x, y, 'Students')
        y -= 24
        c.setFont('Helvetica', 10)
        c.drawString(x, y, 'Name')
        c.drawString(x+220, y, 'Email')
        c.drawString(x+420, y, 'Taken')
        c.drawString(x+480, y, 'Avg')
        y -= 14
        c.line(x, y, width-40, y)
        y -= 14
        for s in students:
            if y < 60:
                c.showPage()
                y = height - 40
            avg = '' if s.get('avg_percent') is None else f"{s.get('avg_percent')}%"
            c.drawString(x, y, str(s.get('name') or ''))
            c.drawString(x+220, y, str(s.get('email') or ''))
            c.drawString(x+420, y, str(s.get('total_taken') or ''))
            c.drawString(x+480, y, avg)
            y -= 16

        c.save()
        bio.seek(0)
        return send_file(bio, as_attachment=True, download_name='students.pdf', mimetype='application/pdf')

    # unknown format -> redirect back
    flash('Unknown export format.', 'error')
    return redirect(url_for('main.teacher_students'))


@main.route('/teacher/student/<int:student_id>')
def teacher_view_student(student_id):
    """Teacher-facing view of a single student's progress.
    Collects the same metrics as the student dashboard/progress endpoints but for the
    specified student_id and renders the partial `view_student.html`.
    """
    # load student
    student = Student.query.filter_by(id=student_id).first()
    if not student:
        flash('Student not found.', 'error')
        return redirect(url_for('main.teacher_students'))

    # recent quizzes attempted by this student
    attempts = (
        QuizAttempt.query.filter_by(student_id=student_id)
        .order_by(QuizAttempt.completed_at.desc())
        .limit(25)
        .all()
    )

    # build list of recent quiz attempts for this student (most recent first)
    # the template will render quiz info from the related Quiz on each attempt
    quizzes = (
        QuizAttempt.query.filter(QuizAttempt.student_id == student_id, QuizAttempt.completed_at != None)
        .order_by(QuizAttempt.completed_at.desc())
        .limit(10)
        .all()
    )

    # aggregate metrics similar to student_dashboard
    quizzes_taken = QuizAttempt.query.filter(QuizAttempt.student_id == student_id, QuizAttempt.completed_at != None).count()

    avg_score = None
    try:
        avg_val = db.session.query(func.avg(QuizAttempt.percent)).filter(QuizAttempt.student_id == student_id, QuizAttempt.completed_at != None, QuizAttempt.percent != None).scalar()
        if avg_val is not None:
            avg_score = round(float(avg_val), 1)
    except Exception:
        avg_score = None

    # streak
    days_streak = 0
    try:
        attempts_dates = (
            QuizAttempt.query.filter(QuizAttempt.student_id == student_id, QuizAttempt.completed_at != None)
            .with_entities(QuizAttempt.completed_at)
            .order_by(QuizAttempt.completed_at.desc())
            .all()
        )
        completed_dates = set()
        for (dt,) in attempts_dates:
            if dt:
                try:
                    completed_dates.add(dt.date())
                except Exception:
                    continue
        today = datetime.utcnow().date()
        cur = today
        streak = 0
        while cur in completed_dates:
            streak += 1
            cur = cur - timedelta(days=1)
        days_streak = streak
    except Exception:
        days_streak = 0

    # ranking among students
    current_student_rank = None
    current_student_avg = None
    try:
        averages = (
            db.session.query(
                Student.id.label('student_id'),
                func.avg(QuizAttempt.percent).label('avg_percent')
            )
            .join(QuizAttempt, QuizAttempt.student_id == Student.id)
            .filter(QuizAttempt.completed_at != None, QuizAttempt.percent != None)
            .group_by(Student.id)
            .order_by(func.avg(QuizAttempt.percent).desc())
            .all()
        )
        last_score = None
        dense_rank = 0
        for row in averages:
            avg = float(row.avg_percent) if row.avg_percent is not None else 0.0
            if last_score is None or avg != last_score:
                dense_rank += 1
            last_score = avg
            if row.student_id == student_id:
                current_student_rank = dense_rank
                current_student_avg = round(avg, 1)
                break
    except Exception:
        current_student_rank = None
        current_student_avg = None

    # today's goal and completed count
    dg = student.daily_goal if student and getattr(student, 'daily_goal', None) is not None else 1
    dc = QuizAttempt.query.filter(QuizAttempt.student_id == student_id, QuizAttempt.completed_at != None, QuizAttempt.completed_at >= (datetime.utcnow() - timedelta(days=1))).count()

    # weekly goal calculation (mirror student_progress)
    try:
        weekly_goal_total = 5
        if student and getattr(student, 'weekly_goal', None) is not None:
            try:
                weekly_goal_total = int(student.weekly_goal)
            except Exception:
                weekly_goal_total = 5
        now = datetime.utcnow()
        week_start = now - timedelta(days=7)
        weekly_goal_completed = QuizAttempt.query.filter(QuizAttempt.student_id == student_id, QuizAttempt.completed_at != None, QuizAttempt.completed_at >= week_start).count()
        weekly_goal_percent = int((weekly_goal_completed / weekly_goal_total) * 100) if weekly_goal_total > 0 else 0
        if weekly_goal_percent > 100:
            weekly_goal_percent = 100
        weekly_goal_message = "Keep going!" if weekly_goal_percent < 100 else "Goal reached \u2014 great job!"
    except Exception:
        weekly_goal_total = 5
        weekly_goal_completed = 0
        weekly_goal_percent = 0
        weekly_goal_message = "Weekly progress unavailable"

    # compute per-subject average percent for this student (strengths/weaknesses)
    subject_strengths = []
    try:
        rows = (
            db.session.query(
                Subject.id.label('subject_id'),
                Subject.name.label('subject_name'),
                func.avg(QuizAttempt.percent).label('avg_percent'),
                func.count(QuizAttempt.id).label('attempts_count')
            )
            .join(Quiz, Quiz.subject_id == Subject.id)
            .join(QuizAttempt, QuizAttempt.quiz_id == Quiz.id)
            .filter(QuizAttempt.student_id == student_id, QuizAttempt.completed_at != None, QuizAttempt.percent != None)
            .group_by(Subject.id)
            .order_by(func.avg(QuizAttempt.percent).desc())
            .all()
        )
        for r in rows:
            subject_strengths.append({
                'id': r.subject_id,
                'name': r.subject_name,
                'avg_percent': round(float(r.avg_percent), 1) if r.avg_percent is not None else 0.0,
                'attempts': int(r.attempts_count),
            })
    except Exception:
        subject_strengths = []

    return render_template('partials/view_student.html', quizzes=quizzes, daily_goal=dg, daily_completed=dc, quizzes_taken=quizzes_taken, avg_score=avg_score, days_streak=days_streak, current_student_rank=current_student_rank, current_student_avg=current_student_avg, student=student, weekly_goal_total=weekly_goal_total, weekly_goal_completed=weekly_goal_completed, weekly_goal_percent=weekly_goal_percent, weekly_goal_message=weekly_goal_message, subject_strengths=subject_strengths)


@main.route('/teacher/student/<int:student_id>/export')
def teacher_export_student_attempts(student_id):
    """Export a single student's recent quiz attempts in csv, txt or docx formats.
    Usage: /teacher/student/<id>/export?format=csv|txt|docx
    """
    fmt = (request.args.get('format') or 'csv').lower()

    # load student and their recent attempts
    student = Student.query.filter_by(id=student_id).first()
    if not student:
        flash('Student not found.', 'error')
        return redirect(url_for('main.teacher_students'))

    attempts = (
        QuizAttempt.query.filter(QuizAttempt.student_id == student_id, QuizAttempt.completed_at != None)
        .order_by(QuizAttempt.completed_at.desc())
        .limit(50)
        .all()
    )

    # prepare rows
    rows = []
    for a in attempts:
        q = a.quiz
        title = q.title if q else f'Quiz #{a.quiz_id}'
        subject = q.subject.name if q and q.subject else ''
        completed = a.completed_at.strftime('%Y-%m-%d %H:%M:%S') if a.completed_at else ''
        score = a.score if a.score is not None else ''
        percent = ('%.1f' % a.percent) if a.percent is not None else ''
        time_taken = f"{int(a.time_taken_seconds // 60):02d}:{int(a.time_taken_seconds % 60):02d}" if a.time_taken_seconds else ''
        rows.append([title, subject, completed, score, percent, time_taken])

    # TXT
    if fmt == 'txt':
        out = io.StringIO()
        out.write('Title\tSubject\tCompleted At\tScore\tPercent\tTime\n')
        for r in rows:
            out.write('\t'.join([str(x) for x in r]) + '\n')
        bio = io.BytesIO()
        bio.write(out.getvalue().encode('utf-8'))
        bio.seek(0)
        filename = f"{(student.name or 'student').replace(' ', '_')}_attempts.txt"
        return send_file(bio, as_attachment=True, download_name=filename, mimetype='text/plain')

    # CSV (Excel friendly)
    if fmt == 'csv':
        out = io.StringIO()
        writer = csv.writer(out)
        writer.writerow(['Title', 'Subject', 'Completed At', 'Score', 'Percent', 'Time'])
        for r in rows:
            writer.writerow(r)
        bio = io.BytesIO()
        bio.write(out.getvalue().encode('utf-8-sig'))
        bio.seek(0)
        filename = f"{(student.name or 'student').replace(' ', '_')}_attempts.csv"
        return send_file(bio, as_attachment=True, download_name=filename, mimetype='text/csv')

    # DOCX (Word) - optional dependency: python-docx
    if fmt == 'docx':
        # Prefer python-docx for a proper .docx when available
        try:
            from docx import Document
            has_docx = True
        except Exception:
            has_docx = False

        if has_docx:
            doc = Document()
            doc.add_heading(f"Quiz Attempts - {student.name}", level=1)
            table = doc.add_table(rows=1, cols=6)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Title'
            hdr_cells[1].text = 'Subject'
            hdr_cells[2].text = 'Completed At'
            hdr_cells[3].text = 'Score'
            hdr_cells[4].text = 'Percent'
            hdr_cells[5].text = 'Time'
            for r in rows:
                row_cells = table.add_row().cells
                for i, val in enumerate(r):
                    row_cells[i].text = str(val or '')

            bio = io.BytesIO()
            doc.save(bio)
            bio.seek(0)
            filename = f"{(student.name or 'student').replace(' ', '_')}_attempts.docx"
            return send_file(bio, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        # Fallback: return a simple HTML document with .doc extension — Word can open this
        html = ['<html><head><meta charset="utf-8" /><title>Quiz Attempts</title></head><body>']
        html.append(f'<h1>Quiz Attempts - {student.name}</h1>')
        html.append('<table border="1" cellspacing="0" cellpadding="4">')
        html.append('<tr><th>Title</th><th>Subject</th><th>Completed At</th><th>Score</th><th>Percent</th><th>Time</th></tr>')
        for r in rows:
            html.append('<tr>' + ''.join(f'<td>{(x or "")}</td>' for x in r) + '</tr>')
        html.append('</table></body></html>')
        bio = io.BytesIO()
        bio.write('\n'.join(html).encode('utf-8'))
        bio.seek(0)
        filename = f"{(student.name or 'student').replace(' ', '_')}_attempts.doc"
        return send_file(bio, as_attachment=True, download_name=filename, mimetype='application/msword')

    flash('Unknown export format.', 'error')
    return redirect(url_for('main.teacher_view_student', student_id=student_id))

@main.route("/teacher/messages")
def teacher_messages():
    # list conversations for this teacher (show student names and unread counts)
    teacher_id = session.get('teacher_id')
    conversations = []
    if teacher_id:
        convs = Conversation.query.filter_by(teacher_id=teacher_id).order_by(Conversation.created_at.desc()).all()
        for c in convs:
            # last message snippet & unread count
            last = None
            unread = 0
            if c.messages:
                last = c.messages[-1]
                unread = Message.query.filter(Message.conversation_id == c.id, Message.sender_role == 'student', Message.read == False).count()
            conversations.append({'id': c.id, 'student': c.student, 'last': last, 'unread': unread})
    # also provide a list of students to start a new conversation
    students = Student.query.order_by(Student.name.asc()).all()
    return render_template("teacher/messages.html", conversations=conversations, students=students)


@main.route("/student/messages")
def student_messages():
    # list conversations for this student
    student_id = session.get('student_id')
    conversations = []
    if student_id:
        convs = Conversation.query.filter_by(student_id=student_id).order_by(Conversation.created_at.desc()).all()
        for c in convs:
            last = None
            unread = 0
            if c.messages:
                last = c.messages[-1]
                unread = Message.query.filter(Message.conversation_id == c.id, Message.sender_role == 'teacher', Message.read == False).count()
            conversations.append({'id': c.id, 'teacher': c.teacher, 'last': last, 'unread': unread})
    # include list of teachers for starting a new conversation
    teachers = Teacher.query.order_by(Teacher.name.asc()).all()
    return render_template('student/messages.html', conversations=conversations, teachers=teachers)


# API: fetch messages for a conversation
@main.route('/api/messages/<int:conversation_id>')
def api_get_messages(conversation_id):
    conv = Conversation.query.filter_by(id=conversation_id).first()
    if not conv:
        return jsonify({'error': 'Conversation not found'}), 404
    msgs = []
    for m in conv.messages:
        msgs.append({'id': m.id, 'sender_role': m.sender_role, 'sender_id': m.sender_id, 'text': m.text, 'created_at': m.created_at.isoformat(), 'read': m.read})
    return jsonify({'conversation_id': conv.id, 'messages': msgs})


# API: send a message (creates conversation if needed)
@main.route('/api/messages/send', methods=['POST'])
def api_send_message():
    data = request.get_json() or {}
    conversation_id = data.get('conversation_id')
    text = data.get('text')
    sender_role = data.get('sender_role')

    if not text or not sender_role:
        return jsonify({'error': 'Missing fields'}), 400

    # determine sender id based on role from session
    sender_id = None
    if sender_role == 'teacher':
        sender_id = session.get('teacher_id')
    elif sender_role == 'student':
        sender_id = session.get('student_id')
    else:
        return jsonify({'error': 'Invalid sender_role'}), 400

    if not sender_id:
        return jsonify({'error': 'Authentication required'}), 401

    conv = None
    if conversation_id:
        conv = Conversation.query.filter_by(id=conversation_id).first()

    # If conversation missing, optionally create when teacher->student or student->teacher
    if not conv:
        # require the other_party id in the payload
        other_id = data.get('other_id')
        if not other_id:
            return jsonify({'error': 'conversation_id missing and other_id not provided to create conversation'}), 400
        if sender_role == 'teacher':
            conv = Conversation(teacher_id=sender_id, student_id=other_id)
        else:
            conv = Conversation(teacher_id=other_id, student_id=sender_id)
        db.session.add(conv)
        db.session.flush()

    try:
        msg = Message(conversation_id=conv.id, sender_role=sender_role, sender_id=sender_id, text=text)
        db.session.add(msg)
        db.session.commit()
        return jsonify({'status': 'ok', 'message_id': msg.id, 'conversation_id': conv.id}), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


# API: mark messages as read in a conversation for the current user
@main.route('/api/messages/<int:conversation_id>/read', methods=['POST'])
def api_mark_read(conversation_id):
    conv = Conversation.query.filter_by(id=conversation_id).first()
    if not conv:
        return jsonify({'error': 'Conversation not found'}), 404
    # determine viewer role
    teacher_id = session.get('teacher_id')
    student_id = session.get('student_id')
    # mark messages sent by the opposite role as read
    try:
        if teacher_id and conv.teacher_id == teacher_id:
            Message.query.filter(Message.conversation_id == conv.id, Message.sender_role == 'student', Message.read == False).update({'read': True})
        elif student_id and conv.student_id == student_id:
            Message.query.filter(Message.conversation_id == conv.id, Message.sender_role == 'teacher', Message.read == False).update({'read': True})
        else:
            return jsonify({'error': 'Not a participant'}), 403
        db.session.commit()
        return jsonify({'status': 'ok'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@main.route('/api/conversations/get_or_create', methods=['POST'])
def api_get_or_create_conversation():
    data = request.get_json() or {}
    student_id = data.get('student_id')

    teacher_id = session.get('teacher_id')
    if not teacher_id:
        return jsonify({'error': 'Authentication required (teacher)'}), 401
    if not student_id:
        return jsonify({'error': 'student_id required'}), 400

    try:
        # check for existing conversation
        conv = Conversation.query.filter_by(teacher_id=teacher_id, student_id=student_id).first()
        if conv:
            return jsonify({'conversation_id': conv.id, 'created': False}), 200

        # create new one
        conv = Conversation(teacher_id=teacher_id, student_id=student_id)
        db.session.add(conv)
        db.session.commit()
        return jsonify({'conversation_id': conv.id, 'created': True}), 201
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@main.route("/teacher/analytics")
def teacher_analytics():
    # Build a simple performance trend: average percent across all students
    # for the last 6 months (including current month). We'll compute month buckets
    # and average the QuizAttempt.percent for completed attempts in each month.
    try:
        now = datetime.utcnow()
        # prepare 6 months range (including current month)
        months = []
        for i in range(5, -1, -1):
            m = (now - timedelta(days=30 * i))
            months.append((m.year, m.month))

        # build month labels and start/end datetimes for each bucket
        buckets = []
        labels = []
        for y, mo in months:
            labels.append(datetime(y, mo, 1).strftime('%b'))
            start = datetime(y, mo, 1)
            # compute end as first day of next month
            if mo == 12:
                end = datetime(y + 1, 1, 1)
            else:
                end = datetime(y, mo + 1, 1)
            buckets.append((start, end))

        avg_percents = []
        for start, end in buckets:
            val = (
                db.session.query(func.avg(QuizAttempt.percent))
                .filter(QuizAttempt.completed_at != None, QuizAttempt.completed_at >= start, QuizAttempt.completed_at < end, QuizAttempt.percent != None)
                .scalar()
            )
            try:
                if val is None:
                    avg_percents.append(None)
                else:
                    avg_percents.append(round(float(val), 1))
            except Exception:
                avg_percents.append(None)

        # Convert averages into SVG coordinates for the existing 600x220 viewBox
        # y-axis: 20..180 maps to 100..0 (percent -> pixel)
        # x positions: spread evenly between 60 and 560
        xs = [60 + i * 100 for i in range(len(avg_percents))]
        points = []
        circles = []
        for i, v in enumerate(avg_percents):
            x = xs[i]
            if v is None:
                # place missing points at bottom (or skip)
                y = 180
            else:
                # percent 100 -> y=20, 0 -> y=180
                y = 20 + (100 - max(0, min(100, v))) * (160 / 100.0)
            points.append((x, int(y)))
            circles.append({'x': x, 'y': int(y), 'label': (str(v) + '%' if v is not None else 'n/a')})

        # build polyline points string skipping trailing None values to avoid drops
        points_str = ' '.join(f"{p[0]},{p[1]}" for p in points)

    except Exception:
        labels = ['Jan','Feb','Mar','Apr','May','Jun']
        points_str = "60,120 160,90 260,80 360,90 460,70 560,80"
        circles = [
            {'x':60,'y':120,'label':'80%'},
            {'x':160,'y':90,'label':'90%'},
            {'x':260,'y':80,'label':'93%'},
            {'x':360,'y':90,'label':'90%'},
            {'x':460,'y':70,'label':'95%'},
            {'x':560,'y':80,'label':'93%'},
        ]

    return render_template("teacher/analytics.html", trend_points=points_str, trend_circles=circles, trend_labels=labels)

@main.route("/teacher/settings")
def teacher_settings():
    teacher_id = session.get('teacher_id')
    teacher = None
    if teacher_id:
        teacher = Teacher.query.filter_by(id=teacher_id).first()
    return render_template("teacher/settings.html", teacher=teacher)


@main.route('/teacher/settings/update_info', methods=['POST'])
def teacher_update_info():
    teacher_id = session.get('teacher_id')
    if not teacher_id:
        flash('Authentication required.', 'error')
        return redirect(url_for('auth.login_teacher'))
    teacher = Teacher.query.filter_by(id=teacher_id).first()
    if not teacher:
        flash('Teacher not found.', 'error')
        return redirect(url_for('auth.login_teacher'))

    name = request.form.get('name')
    email = request.form.get('email')
    if not name or not email:
        flash('Name and email are required.', 'error')
        return redirect(url_for('main.teacher_settings'))

    # check for email collision
    existing = Teacher.query.filter(Teacher.email == email, Teacher.id != teacher.id).first()
    if existing:
        flash('Email already in use by another account.', 'error')
        return redirect(url_for('main.teacher_settings'))

    teacher.name = name
    teacher.email = email
    db.session.commit()
    flash('Profile updated successfully.', 'success')
    return redirect(url_for('main.teacher_settings'))


@main.route('/teacher/settings/update_password', methods=['POST'])
def teacher_update_password():
    teacher_id = session.get('teacher_id')
    if not teacher_id:
        flash('Authentication required.', 'error')
        return redirect(url_for('auth.login_teacher'))
    teacher = Teacher.query.filter_by(id=teacher_id).first()
    if not teacher:
        flash('Teacher not found.', 'error')
        return redirect(url_for('auth.login_teacher'))

    current = request.form.get('current_password')
    new = request.form.get('new_password')
    confirm = request.form.get('confirm_password')
    if not current or not new or not confirm:
        flash('Please fill out all password fields.', 'error')
        return redirect(url_for('main.teacher_settings'))
    if not teacher.check_password(current):
        flash('Current password is incorrect.', 'error')
        return redirect(url_for('main.teacher_settings'))
    if new != confirm:
        flash('New passwords do not match.', 'error')
        return redirect(url_for('main.teacher_settings'))
    teacher.set_password(new)
    db.session.commit()
    flash('Password updated successfully.', 'success')
    return redirect(url_for('main.teacher_settings'))


@main.route('/teacher/settings/delete_account', methods=['POST'])
def teacher_delete_account():
    teacher_id = session.get('teacher_id')
    if not teacher_id:
        flash('Authentication required.', 'error')
        return redirect(url_for('auth.login_teacher'))
    teacher = Teacher.query.filter_by(id=teacher_id).first()
    if not teacher:
        flash('Teacher not found.', 'error')
        return redirect(url_for('auth.login_teacher'))

    # Optional: confirm via password
    password = request.form.get('confirm_password')
    if password and not teacher.check_password(password):
        flash('Password confirmation incorrect.', 'error')
        return redirect(url_for('main.teacher_settings'))

    # delete teacher and logout
    db.session.delete(teacher)
    db.session.commit()
    session.clear()
    flash('Your account has been deleted.', 'success')
    return redirect(url_for('auth.login_teacher'))


# Edit quiz route
@main.route("/teacher/quizzes/edit/<int:quiz_id>", methods=["GET", "POST"])
def teacher_edit_quiz(quiz_id):
    teacher_id = session.get('teacher_id')
    quiz = Quiz.query.filter_by(id=quiz_id, teacher_id=teacher_id).first()
    if not quiz:
        flash("Quiz not found or access denied.", "error")
        return redirect(url_for('main.teacher_quizzes'))

    if request.method == "GET":
        return render_template("teacher/create_quiz.html", quiz=quiz, edit_mode=True)

    # POST: update quiz
    data = request.get_json() or request.form
    quiz.title = data.get('title', quiz.title)
    quiz.description = data.get('description', quiz.description)
    quiz.time_limit = int(data.get('time_limit', quiz.time_limit))
    quiz.difficulty = data.get('difficulty', quiz.difficulty)
    subject_name = data.get('subject')
    if subject_name:
        subject = Subject.query.filter_by(name=subject_name).first()
        if not subject:
            subject = Subject(name=subject_name)
            db.session.add(subject)
            db.session.flush()
        quiz.subject = subject

    # Update questions
    questions = data.get('questions', [])
    # Remove old questions/options
    for q in quiz.questions:
        for opt in q.options:
            db.session.delete(opt)
        db.session.delete(q)
    db.session.flush()

    # Add new questions/options
    for q in questions:
        qtype = q.get('type')
        qtext = q.get('text')
        correct = q.get('correct')
        opts = q.get('options', [])
        question = Question(quiz=quiz, type=qtype, text=qtext, correct_answer=str(correct))
        db.session.add(question)
        db.session.flush()
        if qtype == 'mc':
            for opt in opts:
                key = opt.get('key')
                text = opt.get('text')
                if key and text is not None:
                    option = Option(question=question, key=key, text=text)
                    db.session.add(option)

    db.session.commit()
    if request.is_json:
        return jsonify({"status": "success", "quiz_id": quiz.id}), 200
    else:
        flash("Quiz updated successfully!", "success")
        return redirect(url_for('main.teacher_quizzes'))


@main.route("/student/progress")
def student_progress():
    student_id = session.get('student_id')
    attempts = []
    if student_id:
        # load recent attempts for this student (most recent first)
        attempts = (
            QuizAttempt.query.filter_by(student_id=student_id)
            .order_by(QuizAttempt.completed_at.desc())
            .limit(25)
            .all()
        )
    # Weekly goal calculation: default goal is 5 quizzes per week
    try:
        # default goal is 5, but prefer student's stored weekly_goal when available
        goal = 5
        if student_id:
            s = Student.query.filter_by(id=student_id).first()
            if s and getattr(s, 'weekly_goal', None) is not None:
                try:
                    goal = int(s.weekly_goal)
                except Exception:
                    goal = 5
        now = datetime.utcnow()
        week_start = now - timedelta(days=7)
        completed_this_week = 0
        if student_id:
            completed_this_week = (
                QuizAttempt.query.filter(QuizAttempt.student_id == student_id, QuizAttempt.completed_at != None, QuizAttempt.completed_at >= week_start)
                .count()
            )

        percent = int((completed_this_week / goal) * 100) if goal > 0 else 0
        if percent > 100:
            percent = 100

        weekly_message = "Keep going!" if percent < 100 else "Goal reached — great job!"
    except Exception:
        # fallback values in case of any DB error
        goal = 5
        completed_this_week = 0
        percent = 0
        weekly_message = "Weekly progress unavailable"

    # compute per-subject average percent for this student
    subject_strengths = []
    if student_id:
        try:
            rows = (
                db.session.query(
                    Subject.id.label('subject_id'),
                    Subject.name.label('subject_name'),
                    func.avg(QuizAttempt.percent).label('avg_percent'),
                    func.count(QuizAttempt.id).label('attempts_count')
                )
                .join(Quiz, Quiz.subject_id == Subject.id)
                .join(QuizAttempt, QuizAttempt.quiz_id == Quiz.id)
                .filter(QuizAttempt.student_id == student_id, QuizAttempt.completed_at != None, QuizAttempt.percent != None)
                .group_by(Subject.id)
                .order_by(func.avg(QuizAttempt.percent).desc())
                .all()
            )
            for r in rows:
                subject_strengths.append({
                    'id': r.subject_id,
                    'name': r.subject_name,
                    'avg_percent': round(float(r.avg_percent), 1) if r.avg_percent is not None else 0.0,
                    'attempts': int(r.attempts_count),
                })
        except Exception:
            subject_strengths = []

    return render_template(
        "student/progress.html",
        attempts=attempts,
        weekly_goal_total=goal,
        weekly_goal_completed=completed_this_week,
        weekly_goal_percent=percent,
        weekly_goal_message=weekly_message,
        subject_strengths=subject_strengths,
    )

@main.route("/student/ranking")
def student_ranking():
    # compute average percent score per student across completed attempts
    # and sort highest to lowest. Also determine the logged-in student's rank.
    student_id = session.get('student_id')

    # Build a query that returns student id, name, average percent and count of attempts
    try:
        # join Student and QuizAttempt, aggregate by student
        averages = (
            db.session.query(
                Student.id.label('student_id'),
                Student.name.label('name'),
                func.avg(QuizAttempt.percent).label('avg_percent'),
                func.count(QuizAttempt.id).label('attempts_count'),
            )
            .join(QuizAttempt, QuizAttempt.student_id == Student.id)
            .filter(QuizAttempt.completed_at != None, QuizAttempt.percent != None)
            .group_by(Student.id)
            .order_by(func.avg(QuizAttempt.percent).desc())
            .all()
        )

        # Convert to list of dicts and assign rank numbers (1-based). Handle ties by dense ranking.
        ranked = []
        last_score = None
        current_rank = 0
        dense_rank = 0
        for row in averages:
            avg = float(row.avg_percent) if row.avg_percent is not None else 0.0
            # increment dense rank when score differs
            if last_score is None or avg != last_score:
                dense_rank = dense_rank + 1
            last_score = avg
            ranked.append({
                'rank': dense_rank,
                'student_id': row.student_id,
                'name': row.name,
                'avg_percent': round(avg, 1),
                'attempts_count': int(row.attempts_count),
            })

        # find current student's rank if logged in
        current_student_rank = None
        current_student_avg = None
        if student_id:
            for s in ranked:
                if s['student_id'] == student_id:
                    current_student_rank = s['rank']
                    current_student_avg = s['avg_percent']
                    break

    except Exception:
        ranked = []
        current_student_rank = None
        current_student_avg = None

    # Try to fetch current student's name and total completed attempts for display
    current_student_name = None
    current_student_attempts = 0
    try:
        if student_id:
            s = Student.query.filter_by(id=student_id).first()
            if s:
                current_student_name = s.name
                current_student_attempts = (
                    QuizAttempt.query.filter(QuizAttempt.student_id == student_id, QuizAttempt.completed_at != None).count()
                )
    except Exception:
        current_student_name = current_student_name
        current_student_attempts = current_student_attempts

    # pass top 5 and full ranked list to template
    top_students = ranked[:5]
    return render_template(
        "student/ranking.html",
        top_students=top_students,
        ranked_students=ranked,
        current_student_rank=current_student_rank,
        current_student_avg=current_student_avg,
        current_student_id=student_id,
        current_student_name=current_student_name,
        current_student_attempts=current_student_attempts,
    )

# Student settings (mirror teacher settings)
@main.route("/student/settings")
def student_settings():
    student_id = session.get('student_id')
    student = None
    if student_id:
        student = Student.query.filter_by(id=student_id).first()
    return render_template("student/settings.html", student=student)


@main.route('/student/settings/update_info', methods=['POST'])
def student_update_info():
    student_id = session.get('student_id')
    if not student_id:
        flash('Authentication required.', 'error')
        return redirect(url_for('auth.login_student'))
    student = Student.query.filter_by(id=student_id).first()
    if not student:
        flash('Student not found.', 'error')
        return redirect(url_for('auth.login_student'))

    name = request.form.get('name')
    email = request.form.get('email')
    if not name or not email:
        flash('Name and email are required.', 'error')
        return redirect(url_for('main.student_settings'))

    # check for email collision
    existing = Student.query.filter(Student.email == email, Student.id != student.id).first()
    if existing:
        flash('Email already in use by another account.', 'error')
        return redirect(url_for('main.student_settings'))

    student.name = name
    student.email = email
    db.session.commit()
    flash('Profile updated successfully.', 'success')
    return redirect(url_for('main.student_settings'))


@main.route('/student/settings/update_password', methods=['POST'])
def student_update_password():
    student_id = session.get('student_id')
    if not student_id:
        flash('Authentication required.', 'error')
        return redirect(url_for('auth.login_student'))
    student = Student.query.filter_by(id=student_id).first()
    if not student:
        flash('Student not found.', 'error')
        return redirect(url_for('auth.login_student'))

    current = request.form.get('current_password')
    new = request.form.get('new_password')
    confirm = request.form.get('confirm_password')
    if not current or not new or not confirm:
        flash('Please fill out all password fields.', 'error')
        return redirect(url_for('main.student_settings'))
    if not student.check_password(current):
        flash('Current password is incorrect.', 'error')
        return redirect(url_for('main.student_settings'))
    if new != confirm:
        flash('New passwords do not match.', 'error')
        return redirect(url_for('main.student_settings'))
    student.set_password(new)
    db.session.commit()
    flash('Password updated successfully.', 'success')
    return redirect(url_for('main.student_settings'))


@main.route('/student/settings/delete_account', methods=['POST'])
def student_delete_account():
    student_id = session.get('student_id')
    if not student_id:
        flash('Authentication required.', 'error')
        return redirect(url_for('auth.login_student'))
    student = Student.query.filter_by(id=student_id).first()
    if not student:
        flash('Student not found.', 'error')
        return redirect(url_for('auth.login_student'))

    # Optional: confirm via password
    password = request.form.get('confirm_password')
    if password and not student.check_password(password):
        flash('Password confirmation incorrect.', 'error')
        return redirect(url_for('main.student_settings'))

    # delete student and logout
    db.session.delete(student)
    db.session.commit()
    session.clear()
    flash('Your account has been deleted.', 'success')
    return redirect(url_for('auth.login_student'))


@main.route('/student/settings/update_goals', methods=['POST'])
def student_update_goals():
    student_id = session.get('student_id')
    if not student_id:
        flash('Authentication required.', 'error')
        return redirect(url_for('auth.login_student'))
    student = Student.query.filter_by(id=student_id).first()
    if not student:
        flash('Student not found.', 'error')
        return redirect(url_for('auth.login_student'))

    # parse and validate goals
    daily = request.form.get('daily_goal')
    weekly = request.form.get('weekly_goal')
    try:
        daily_val = int(daily) if daily is not None and str(daily).strip() != '' else None
    except ValueError:
        flash('Daily goal must be a number.', 'error')
        return redirect(url_for('main.student_settings'))
    try:
        weekly_val = int(weekly) if weekly is not None and str(weekly).strip() != '' else None
    except ValueError:
        flash('Weekly goal must be a number.', 'error')
        return redirect(url_for('main.student_settings'))

    if daily_val is not None:
        if daily_val < 0:
            flash('Daily goal must be zero or positive.', 'error')
            return redirect(url_for('main.student_settings'))
        student.daily_goal = daily_val
    if weekly_val is not None:
        if weekly_val < 0:
            flash('Weekly goal must be zero or positive.', 'error')
            return redirect(url_for('main.student_settings'))
        student.weekly_goal = weekly_val

    db.session.commit()
    flash('Goals updated successfully.', 'success')
    return redirect(url_for('main.student_settings'))
